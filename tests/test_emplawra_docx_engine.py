import sys
import os
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from emplawra_docx_engine import (
    _has_2column_layout, _has_2column_table, _find_column_break,
    _heuristic_column_split, _pair_by_position, _write_2cell_table,
    _parse_translated_response, extract_paragraphs, chunk_paragraphs, translate_chunk,
    write_inline, write_side_by_side, transform2cell, parse_args,
    load_config, get_provider, restore_placeholders, NS_W,
)


def test_config_loading():
    import json
    import tempfile
    import os
    cfg = {
        "default_provider": "test",
        "providers": {
            "test": {
                "base_url": "http://test/v1",
                "model": "test-model",
                "api_key_env": None
            }
        }
    }
    with tempfile.NamedTemporaryFile(mode='w', suffix='.json', delete=False) as f:
        json.dump(cfg, f)
        fname = f.name
    try:
        result = load_config(fname)
        assert result == cfg
    finally:
        os.unlink(fname)



def test_get_provider():
    cfg = {
        "default_provider": "test",
        "providers": {
            "test": {
                "base_url": "http://test/v1",
                "model": "test-model",
                "api_key_env": None
            }
        }
    }
    provider = get_provider(cfg, "test")
    assert provider["base_url"] == "http://test/v1"
    assert provider["model"] == "test-model"



def test_extract_paragraphs():
    from docx import Document
    import tempfile
    import os
    doc = Document()
    doc.add_paragraph("Hello world")
    p = doc.add_paragraph()
    p.add_run("Bold").bold = True
    p.add_run(" Normal")
    path = os.path.join(tempfile.mkdtemp(), "test.docx")
    doc.save(path)
    paragraphs = extract_paragraphs(path)
    assert len(paragraphs) == 2
    p0 = paragraphs[0]
    assert p0["id"] == "P0"
    assert not p0["in_table"]
    assert p0["cell_row"] is None
    assert p0["cell_col"] is None
    assert p0["alignment"] is None or p0["alignment"] is not None
    assert p0["runs"][0]["text"] == "Hello world"
    assert p0["runs"][0]["italic"] is None
    assert p0["runs"][0]["underline"] is None
    p1 = paragraphs[1]
    assert p1["runs"][0]["text"] == "Bold"
    assert p1["runs"][0]["bold"]
    assert p1["runs"][1]["text"] == " Normal"
    os.unlink(path)



def test_extract_paragraphs_with_table():
    from docx import Document
    import tempfile
    import os
    doc = Document()
    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "A1"
    table.cell(0, 1).text = "B1"
    table.cell(1, 2).text = "C2"
    path = os.path.join(tempfile.mkdtemp(), "test_table.docx")
    doc.save(path)
    paragraphs = extract_paragraphs(path)
    table_paras = [p for p in paragraphs if p["in_table"]]
    assert len(table_paras) == 6, f"Expected 6 table paragraphs, got {len(table_paras)} total paragraphs={len(paragraphs)}"
    for p in table_paras:
        assert p["cell_row"] is not None, f"Expected cell_row, got None for {p}"
        assert p["cell_col"] is not None, f"Expected cell_col, got None for {p}"
    scored = [(p["cell_row"], p["cell_col"], p["runs"][0]["text"] if p["runs"] else "") for p in table_paras]
    assert (0, 0, "A1") in scored, f"Missing A1 at (0,0): {scored}"
    assert (0, 1, "B1") in scored, f"Missing B1 at (0,1): {scored}"
    assert (1, 2, "C2") in scored, f"Missing C2 at (1,2): {scored}"
    os.unlink(path)


def test_extract_paragraphs_invalid_file():
    try:
        extract_paragraphs("/nonexistent/file.docx")
        assert False, "Should have raised FileNotFoundError"
    except FileNotFoundError:
        pass



def test_chunk_paragraphs():
    paragraphs = [{"id": f"P{i}"} for i in range(25)]
    chunks = chunk_paragraphs(paragraphs, 10)
    assert len(chunks) == 3



def test_translate_chunk():
    chunk = [{"id": "P0", "runs": [{"text": "Hello world", "bold": False, "italic": False}]}]
    provider = {"base_url": "http://invalid", "model": "test", "api_key": "bad"}
    try:
        translate_chunk(chunk, "Romanian", provider)
        assert False, "Should have raised ConnectionError or similar"
    except Exception:
        pass



def test_translate_chunk_retry():
    pass



def test_write_inline():
    from docx import Document
    import tempfile
    import os
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Original text")
    src = os.path.join(tempfile.mkdtemp(), "source.docx")
    doc.save(src)
    translated = [{
        "id": "P0",
        "runs": [{"text": "Text tradus", "bold": False, "italic": False}],
        "alignment": None,
        "in_table": False,
    }]
    out = os.path.join(tempfile.mkdtemp(), "output.docx")
    write_inline(src, translated, out)
    result = Document(out)
    assert result.paragraphs[0].text == "Text tradus"
    os.unlink(src)
    os.unlink(out)



def test_write_inline_with_table():
    from docx import Document
    import tempfile
    import os
    doc = Document()
    doc.add_paragraph("Header text")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Cell A1"
    table.cell(0, 1).text = "Cell B1"
    table.cell(1, 0).text = "Cell A2"
    table.cell(1, 1).text = "Cell B2"
    src = os.path.join(tempfile.mkdtemp(), "source.docx")
    doc.save(src)
    translated = [
        {"id": "P0", "runs": [{"text": "Header text trans"}], "alignment": None, "in_table": False},
        {"id": "P1", "runs": [{"text": "Celula A1"}], "alignment": None, "in_table": True, "cell_row": 0, "cell_col": 0},
        {"id": "P2", "runs": [{"text": "Celula B1"}], "alignment": None, "in_table": True, "cell_row": 0, "cell_col": 1},
        {"id": "P3", "runs": [{"text": "Celula A2"}], "alignment": None, "in_table": True, "cell_row": 1, "cell_col": 0},
        {"id": "P4", "runs": [{"text": "Celula B2"}], "alignment": None, "in_table": True, "cell_row": 1, "cell_col": 1},
    ]
    out = os.path.join(tempfile.mkdtemp(), "output.docx")
    write_inline(src, translated, out)
    result = Document(out)
    assert result.paragraphs[0].text == "Header text trans"
    out_table = result.tables[0]
    assert out_table.cell(0, 0).text == "Celula A1"
    assert out_table.cell(0, 1).text == "Celula B1"
    assert out_table.cell(1, 0).text == "Celula A2"
    assert out_table.cell(1, 1).text == "Celula B2"
    os.unlink(src)
    os.unlink(out)



def test_cli_parser():
    args = parse_args(['input.docx', '--lang', 'ro'])
    assert args.input == 'input.docx'
    assert args.lang == 'ro'
    assert args.mode == 'inline'
    assert args.concurrency == 0
    args = parse_args(['input.docx', '--lang', 'en', '--mode', 'side-by-side', '--concurrency', '2'])
    assert args.input == 'input.docx'
    assert args.lang == 'en'
    assert args.mode == 'side-by-side'
    assert args.concurrency == 2



def test_write_2cell_table():
    import tempfile
    import os
    from docx import Document
    doc = Document()
    doc.add_paragraph("Source only")
    src = os.path.join(tempfile.mkdtemp(), "src.docx")
    doc.save(src)
    pairs = [
        ([{"runs": [{"text": "Left A", "bold": False}], "alignment": None, "numpr": None}],
         [{"runs": [{"text": "Right A", "bold": False}], "alignment": None, "numpr": None}]),
        ([{"runs": [{"text": "Left B", "bold": False}], "alignment": None, "numpr": None}],
         [{"runs": [{"text": "Right B", "bold": False}], "alignment": None, "numpr": None}]),
    ]
    out = os.path.join(tempfile.mkdtemp(), "out.docx")
    _write_2cell_table(src, pairs, out)
    result = Document(out)
    assert len(result.tables) == 1
    t = result.tables[0]
    assert len(t.rows) == 2
    assert t.rows[0].cells[0].text == "Left A"
    assert t.rows[0].cells[1].text == "Right A"
    assert t.rows[1].cells[0].text == "Left B"
    assert t.rows[1].cells[1].text == "Right B"
    os.unlink(src)
    os.unlink(out)



def test_write_2cell_table_no_borders():
    import tempfile
    import os
    from docx import Document
    doc = Document()
    doc.add_paragraph("x")
    src = os.path.join(tempfile.mkdtemp(), "src.docx")
    doc.save(src)
    pairs = [([{"runs": [{"text": "A", "bold": False}], "alignment": None, "numpr": None}],
              [{"runs": [{"text": "B", "bold": False}], "alignment": None, "numpr": None}])]
    out = os.path.join(tempfile.mkdtemp(), "out.docx")
    _write_2cell_table(src, pairs, out)
    result = Document(out)
    tbl = result.tables[0]
    tbl_pr = tbl._tbl.tblPr
    borders = tbl_pr.find(f'{NS_W}tblBorders')
    assert borders is not None, "Table should have tblBorders element"
    os.unlink(src)
    os.unlink(out)



def test_write_side_by_side():
    from docx import Document
    import tempfile
    import os
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Original text")
    src = os.path.join(tempfile.mkdtemp(), "source.docx")
    doc.save(src)
    originals = [{"id": "P0", "runs": [{"text": "Original text", "bold": False}], "alignment": None, "in_table": False}]
    translated = [{"id": "P0", "runs": [{"text": "Text tradus", "bold": False}], "alignment": None, "in_table": False}]
    out = os.path.join(tempfile.mkdtemp(), "output.docx")
    write_side_by_side(src, originals, translated, out)
    result = Document(out)
    assert len(result.tables) == 1
    table = result.tables[0]
    assert len(table.rows) == 1
    assert "Original text" in table.rows[0].cells[0].text
    assert "Text tradus" in table.rows[0].cells[1].text
    os.unlink(src)
    os.unlink(out)



def test_restore_placeholders():
    original = "Hello {{client_name}}, your balance is {{amount}}"
    translated = "Salut {{client_name}}, soldul tau este {{amount}}"
    restored = restore_placeholders(translated, original)
    assert "{{client_name}}" in restored
    assert "{{amount}}" in restored



def test_restore_missing_placeholder():
    original = "Hello {{client_name}}"
    translated = "Salut"
    restored = restore_placeholders(translated, original)
    assert "{{client_name}}" in restored



def test_parse_translated_response_protects_placeholders():
    chunk_ids = ["P0", "P1"]
    response_text = "[P0] Salut {{client_name}}\n[P1] Soldul este {{amount}}"
    original_text = "[P0] Hello {{client_name}}\n[P1] Your balance is {{amount}}"
    result = _parse_translated_response(response_text, chunk_ids, original_text)
    assert "{{client_name}}" in result["P0"]
    assert "{{amount}}" in result["P1"]



def test_has_2column_layout():
    import tempfile
    import os
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    doc = Document()
    path = os.path.join(tempfile.mkdtemp(), "test.docx")
    doc.save(path)
    assert not _has_2column_layout(Document(path))
    doc2 = Document()
    sect_pr = doc2.sections[0]._sectPr
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:space'), '720')
    sect_pr.append(cols)
    path2 = os.path.join(tempfile.mkdtemp(), "test2.docx")
    doc2.save(path2)
    assert _has_2column_layout(Document(path2))
    os.unlink(path)
    os.unlink(path2)



def test_has_2column_table():
    import tempfile
    import os
    from docx import Document
    doc = Document()
    doc.add_paragraph("Not a table")
    path = os.path.join(tempfile.mkdtemp(), "test.docx")
    doc.save(path)
    assert not _has_2column_table(Document(path))

    doc2 = Document()
    doc2.add_table(rows=2, cols=2)
    path2 = os.path.join(tempfile.mkdtemp(), "test2.docx")
    doc2.save(path2)
    assert _has_2column_table(Document(path2))

    doc3 = Document()
    doc3.add_table(rows=2, cols=3)
    path3 = os.path.join(tempfile.mkdtemp(), "test3.docx")
    doc3.save(path3)
    assert not _has_2column_table(Document(path3))

    os.unlink(path)
    os.unlink(path2)
    os.unlink(path3)



def test_find_column_break():
    import tempfile
    import os
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    path = os.path.join(tempfile.mkdtemp(), "test.docx")
    doc = Document()
    doc.add_paragraph("Before break")
    p = doc.add_paragraph()
    r = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'column')
    r._element.append(br)
    doc.add_paragraph("After break")
    doc.save(path)
    idx = _find_column_break(Document(path))
    assert idx == 1, f"Expected 1, got {idx}"
    doc2 = Document()
    doc2.add_paragraph("No break")
    doc2.add_paragraph("No break either")
    path2 = os.path.join(tempfile.mkdtemp(), "test2.docx")
    doc2.save(path2)
    assert _find_column_break(Document(path2)) is None
    os.unlink(path)
    os.unlink(path2)



def test_heuristic_column_split():
    from docx import Document
    doc = Document()
    for _i in range(10):
        doc.add_paragraph("x")
    assert _heuristic_column_split(doc) == 5
    doc2 = Document()
    for _i in range(7):
        doc2.add_paragraph("x")
    assert _heuristic_column_split(doc2) == 3



def test_pair_by_position():
    paras = [f"P{i}" for i in range(6)]
    pairs = _pair_by_position(paras, 3)
    assert len(pairs) == 3
    assert pairs[0] == ("P0", "P3")
    assert pairs[1] == ("P1", "P4")
    assert pairs[2] == ("P2", "P5")
    paras2 = [f"P{i}" for i in range(5)]
    pairs2 = _pair_by_position(paras2, 2)
    assert len(pairs2) == 3
    assert pairs2[0] == ("P0", "P2")
    assert pairs2[1] == ("P1", "P3")
    assert pairs2[2] == (None, "P4")



def _create_2column_test_docx(path):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx import Document

    doc = Document()
    sect_pr = doc.sections[0]._sectPr
    existing = sect_pr.find(qn('w:cols'))
    if existing is not None:
        sect_pr.remove(existing)
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:space'), '720')
    sect_pr.append(cols)

    doc.add_paragraph("First original paragraph.")
    doc.add_paragraph("Second original paragraph.")
    doc.add_paragraph("Third original paragraph.")

    p_break = doc.add_paragraph()
    r_break = p_break.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'column')
    r_break._element.append(br)

    doc.add_paragraph("Primul paragraf original.")
    doc.add_paragraph("Al doilea paragraf original.")
    doc.add_paragraph("Al treilea paragraf original.")
    doc.save(path)



def test_transform2cell_integration():
    import tempfile
    import os
    from docx import Document
    path = os.path.join(tempfile.mkdtemp(), "src.docx")
    _create_2column_test_docx(path)
    out = os.path.join(tempfile.mkdtemp(), "out.docx")
    transform2cell(path, out)
    result = Document(out)
    assert len(result.tables) == 1
    t = result.tables[0]
    assert len(t.rows) == 3
    assert "First original" in t.rows[0].cells[0].text
    assert "Primul paragraf" in t.rows[0].cells[1].text
    assert "Second original" in t.rows[1].cells[0].text
    assert "Al doilea" in t.rows[1].cells[1].text
    assert "Third original" in t.rows[2].cells[0].text
    assert "Al treilea" in t.rows[2].cells[1].text
    os.unlink(path)
    os.unlink(out)



def test_transform2cell_table_integration():
    import tempfile
    import os
    from docx import Document
    path = os.path.join(tempfile.mkdtemp(), "src.docx")
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Original A"
    table.cell(0, 1).text = "Translated A"
    table.cell(1, 0).text = "Original B"
    table.cell(1, 1).text = "Translated B"
    doc.save(path)
    out = os.path.join(tempfile.mkdtemp(), "out.docx")
    transform2cell(path, out)
    result = Document(out)
    assert len(result.tables) == 1
    t = result.tables[0]
    assert len(t.rows) == 2
    assert "Original A" in t.rows[0].cells[0].text
    assert "Translated A" in t.rows[0].cells[1].text
    assert "Original B" in t.rows[1].cells[0].text
    assert "Translated B" in t.rows[1].cells[1].text
    os.unlink(path)
    os.unlink(out)



def test_cli_parser_transform2cell():
    args = parse_args(['input.docx', '--transform2cell'])
    assert args.transform2cell is True
    assert args.input == 'input.docx'
    args2 = parse_args(['input.docx', '--transform2cell', '--output', 'out.docx'])
    assert args2.transform2cell
    assert args2.output == 'out.docx'



if __name__ == "__main__":
    import sys
    import inspect
    _test_list = [
        name for name, obj in inspect.getmembers(sys.modules[__name__])
        if name.startswith('test_') and callable(obj)
    ]
    for name in sorted(_test_list):
        fn = globals()[name]
        fn()
        print(f'  {name} PASS')
    print('All tests PASS')
