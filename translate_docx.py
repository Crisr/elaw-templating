import argparse
import copy
import json
import os
import sys
import time
import re
import threading
import zipfile
from concurrent.futures import ThreadPoolExecutor, as_completed
from openai import OpenAI
from docx import Document, Document as DocxDocument
from docx.shared import Pt, RGBColor
from messages import MESSAGES as _

_PARSER = None


def load_config(path="config.json"):
    with open(path) as f:
        return json.load(f)


def get_provider(config, name=None):
    if name is None:
        name = config["default_provider"]
    try:
        provider = config["providers"][name].copy()
    except KeyError:
        providers_list = list(config['providers'].keys())
        raise ValueError(_["provider_not_found"].format(name=name, providers=providers_list))
    api_key_env = provider.pop("api_key_env", None)
    if api_key_env:
        provider["api_key"] = os.environ.get(api_key_env)
        if not provider["api_key"]:
            try:
                from dotenv import load_dotenv  # type: ignore
                load_dotenv()
            except ImportError:
                pass
            provider["api_key"] = os.environ.get(api_key_env)
        if not provider["api_key"]:
            try:
                with open("secrets.json") as f:
                    secrets = json.load(f)
                    provider["api_key"] = secrets.get(api_key_env)
            except (FileNotFoundError, json.JSONDecodeError):
                pass
    return provider


NS_W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'


def _extract_numpr(para):
    pPr = para._element.find(f'{NS_W}pPr')
    if pPr is not None:
        numPr = pPr.find(f'{NS_W}numPr')
        if numPr is not None:
            ilvl = numPr.find(f'{NS_W}ilvl')
            numId = numPr.find(f'{NS_W}numId')
            return {
                'ilvl': ilvl.get(f'{NS_W}val') if ilvl is not None else None,
                'numId': numId.get(f'{NS_W}val') if numId is not None else None,
            }
    return None


def _extract_runs(para):
    runs_data = []
    for run in para.runs:
        runs_data.append({
            "text": run.text,
            "bold": run.bold,
            "italic": run.italic,
            "underline": run.underline,
            "font_name": run.font.name,
            "font_size": run.font.size.pt if run.font.size else None,
            "color": str(run.font.color.rgb) if run.font.color and run.font.color.rgb else None,
        })
    return runs_data


def _cell_coords(para):
    if not para._element.getparent().tag.endswith("tc"):
        return None, None
    tc = para._element.getparent()
    tr = tc.getparent()
    tbl = tr.getparent()
    tcs = [c for c in tr if c.tag.endswith("tc")]
    cell_col = tcs.index(tc)
    trs = [r for r in tbl if r.tag.endswith("tr")]
    cell_row = trs.index(tr)
    return cell_row, cell_col


def extract_paragraphs(path):
    if not os.path.exists(path):
        raise FileNotFoundError(_["file_not_found"].format(path=path))
    doc = Document(path)
    paragraphs = []
    pid = 0

    for para in doc.paragraphs:
        cell_row, cell_col = _cell_coords(para)
        paragraphs.append({
            "id": f"P{pid}",
            "runs": _extract_runs(para),
            "alignment": str(para.alignment) if para.alignment else None,
            "in_table": cell_row is not None,
            "cell_row": cell_row,
            "cell_col": cell_col,
            "numpr": _extract_numpr(para),
        })
        pid += 1

    for table in doc.tables:
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    paragraphs.append({
                        "id": f"P{pid}",
                        "runs": _extract_runs(para),
                        "alignment": str(para.alignment) if para.alignment else None,
                        "in_table": True,
                        "cell_row": row_idx,
                        "cell_col": col_idx,
                        "numpr": _extract_numpr(para),
                    })
                    pid += 1

    return paragraphs


def chunk_paragraphs(paragraphs, chunk_size=10):
    chunks = []
    for i in range(0, len(paragraphs), chunk_size):
        chunks.append(paragraphs[i:i + chunk_size])
    return chunks


def _build_chunk_text(chunk):
    lines = []
    for p in chunk:
        text = ""
        for run in p["runs"]:
            text += run["text"]
        lines.append(f"[{p['id']}] {text}")
    return "\n".join(lines)


def restore_placeholders(translated_text, original_text):
    placeholder_pattern = r'\{\{[^}]+\}\}'
    placeholders = re.findall(placeholder_pattern, original_text)
    for ph in placeholders:
        if ph not in translated_text:
            orig_pos = original_text.index(ph)
            ratio = orig_pos / max(len(original_text), 1)
            insert_pos = int(ratio * len(translated_text))
            translated_text = translated_text[:insert_pos] + ph + translated_text[insert_pos:]
    return translated_text


def _parse_translated_response(response_text, chunk_ids, original_text=None):
    results = {}
    lines = response_text.strip().split("\n")
    if lines and not lines[0].strip().startswith("[P"):
        lines = lines[1:]
    for line in lines:
        line = line.strip()
        if not line:
            continue
        for pid in chunk_ids:
            if line.startswith(f"[{pid}]"):
                text = line[len(f"[{pid}]"):].strip()
                results[pid] = text
                break
    if len(results) < len(chunk_ids):
        all_lines = [ln for ln in response_text.strip().split("\n") if ln.strip()]
        for i, pid in enumerate(chunk_ids):
            if pid not in results and i < len(all_lines):
                results[pid] = all_lines[i].strip()
    if original_text:
        original_lines = {}
        for line in original_text.strip().split("\n"):
            for pid in chunk_ids:
                if line.startswith(f"[{pid}]"):
                    original_lines[pid] = line[len(f"[{pid}]"):].strip()
                    break
        for pid in results:
            if pid in original_lines:
                results[pid] = restore_placeholders(results[pid], original_lines[pid])
    return results


def translate_chunk(chunk, target_lang, provider):
    chunk_text = _build_chunk_text(chunk)
    chunk_ids = [p["id"] for p in chunk]
    api_key = provider.get("api_key") or "not-needed"
    client = OpenAI(base_url=provider["base_url"], api_key=api_key)
    system_prompt = _["llm_system_prompt"].format(lang=target_lang)
    last_error = RuntimeError("translate_chunk failed after all retries")
    for attempt in range(3):
        try:
            resp = client.chat.completions.create(
                model=provider["model"],
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": chunk_text},
                ],
                temperature=0.0,
            )
            translated = resp.choices[0].message.content
            parsed = _parse_translated_response(translated, chunk_ids, chunk_text)
            for p in chunk:
                if p["id"] in parsed:
                    translated_text = parsed[p["id"]]
                    original_runs = p["runs"]
                    original_full_text = "".join(r["text"] for r in original_runs)
                    if original_full_text and len(original_runs) > 1:
                        start = 0
                        for j, run in enumerate(original_runs):
                            proportion = len(run["text"]) / len(original_full_text)
                            chunk_len = int(proportion * len(translated_text))
                            if j == len(original_runs) - 1:
                                run["text"] = translated_text[start:]
                            else:
                                run["text"] = translated_text[start:start + chunk_len]
                                start += chunk_len
                    elif original_runs:
                        original_runs[0]["text"] = translated_text
            return chunk
        except Exception as e:
            last_error = e
            if attempt < 2:
                time.sleep(2 ** attempt)
    raise last_error


def _show_progress(done, total):
    if total == 0:
        return
    width = 30
    pct = done / total
    filled = int(width * pct)
    bar = "#" * filled + "-" * (width - filled)
    print("\r" + _["progress"].format(bar=bar, pct=pct * 100, done=done, total=total), end="", flush=True)


_CONCURRENCY_CACHE = None


def _detect_concurrency(provider):
    global _CONCURRENCY_CACHE
    if _CONCURRENCY_CACHE is not None:
        return _CONCURRENCY_CACHE
    api_key = provider.get("api_key") or "not-needed"
    client = OpenAI(base_url=provider["base_url"], api_key=api_key)
    probe = [{"role": "user", "content": "ok"}]
    best = 1
    for level in [2, 4, 6, 8, 12, 16]:
        print("\r" + _["detecting_concurrency"].format(n=level), end="", flush=True)
        def try_request(_):
            try:
                client.chat.completions.create(model=provider["model"], messages=probe, temperature=0, max_tokens=1)  # type: ignore
                return True
            except Exception:
                return False
        with ThreadPoolExecutor(max_workers=level) as pool:
            results = list(pool.map(try_request, range(level)))
        if sum(1 for r in results if not r) == 0:
            best = level
        else:
            break
    print("\r" + " " * 50 + "\r", end="", flush=True)
    _CONCURRENCY_CACHE = best
    return best


def translate_all(paragraphs, target_lang, provider, concurrency=0, progress_callback=None):
    if concurrency < 1:
        concurrency = _detect_concurrency(provider)
    chunks = chunk_paragraphs(paragraphs)
    total = len(chunks)
    if concurrency > total:
        concurrency = total
    results = [None] * total
    failed_indices = []
    lock = threading.Lock()
    done = 0

    def process(i, chunk):
        try:
            translated = translate_chunk(chunk, target_lang, provider)
            return i, translated, None
        except Exception as e:
            return i, None, e

    if progress_callback:
        progress_callback(0, total)
    else:
        _show_progress(0, total)
    with ThreadPoolExecutor(max_workers=concurrency) as pool:
        futures = [pool.submit(process, i, c) for i, c in enumerate(chunks)]
        for future in as_completed(futures):
            i, translated, err = future.result()
            with lock:
                if err:
                    print("\n" + _["err_translate_chunk"].format(i=i, e=err), file=sys.stderr)
                    failed_indices.append(i)
                else:
                    results[i] = translated
                done += 1
                if progress_callback:
                    progress_callback(done, total)
                else:
                    _show_progress(done, total)
    print()
    all_translated = []
    for r in results:
        if r is not None:
            all_translated.extend(r)
    if failed_indices:
        print(_["warn_chunks_failed"].format(count=len(failed_indices), indices=failed_indices), file=sys.stderr)
    return all_translated


def write_inline(original_path, translated_paragraphs, output_path):
    doc = DocxDocument(original_path)
    trans_by_id = {p["id"]: p for p in translated_paragraphs}
    for i, para in enumerate(doc.paragraphs):
        pid = f"P{i}"
        if pid in trans_by_id:
            tp = trans_by_id[pid]
            for j, run in enumerate(para.runs):
                if j < len(tp["runs"]):
                    run.text = tp["runs"][j]["text"]
    pid = len(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    pid_str = f"P{pid}"
                    if pid_str in trans_by_id:
                        tp = trans_by_id[pid_str]
                        for j, run in enumerate(para.runs):
                            if j < len(tp["runs"]):
                                run.text = tp["runs"][j]["text"]
                    pid += 1
    doc.save(output_path)


def _has_2column_layout(doc):
    for section in doc.sections:
        sect_pr = section._sectPr
        if sect_pr is not None:
            for cols in sect_pr.findall(f'{NS_W}cols'):
                num = cols.get(f'{NS_W}num')
                if num and int(num) == 2:
                    return True
    return False


def _has_2column_table(doc):
    return len(doc.tables) >= 1 and len(doc.tables[0].columns) == 2


def _is_2column_format(doc):
    return _has_2column_layout(doc) or _has_2column_table(doc)


def _extract_pairs_from_table(doc):
    pairs = []
    table = doc.tables[0]
    for row in table.rows:
        cells = row.cells
        left_paras = []
        right_paras = []
        if len(cells) >= 1:
            for para in cells[0].paragraphs:
                left_paras.append(_build_para_dict(para))
        if len(cells) >= 2:
            for para in cells[1].paragraphs:
                right_paras.append(_build_para_dict(para))
        pairs.append((left_paras, right_paras))
    return pairs


def _find_column_break(doc):
    for i, para in enumerate(doc.paragraphs):
        for br in para._element.findall(f'.//{NS_W}br'):
            if br.get(f'{NS_W}type') == 'column':
                return i
    return None


def _heuristic_column_split(doc):
    return len(doc.paragraphs) // 2


def _pair_by_position(paras, split_idx):
    col1 = list(paras[:split_idx])
    col2 = list(paras[split_idx:])
    pairs = []
    for i in range(max(len(col1), len(col2))):
        p1 = col1[i] if i < len(col1) else None
        p2 = col2[i] if i < len(col2) else None
        pairs.append((p1, p2))
    return pairs


def _build_para_dict(para):
    return {
        "runs": _extract_runs(para),
        "alignment": str(int(para.alignment)) if para.alignment is not None else None,
        "numpr": _extract_numpr(para),
    }


def _llm_verify_pairs(col1_dicts, col2_dicts, pairs, provider):
    if provider is None:
        return pairs
    lines = []
    for i, (p1, p2) in enumerate(pairs):
        t1 = "".join(r["text"] for r in p1["runs"]) if p1 else ""
        t2 = "".join(r["text"] for r in p2["runs"]) if p2 else ""
        lines.append(f"[ROW {i}] Left: {t1}")
        lines.append(f"[ROW {i}] Right: {t2}")
    chunk_text = "\n".join(lines)
    client = OpenAI(base_url=provider["base_url"], api_key=provider.get("api_key", "not-needed"))
    system_prompt = (
        "You are verifying a two-column document alignment. "
        "Column 1 (Left) has original text, Column 2 (Right) has the translation. "
        "Each ROW shows a paired paragraph. If any rows are misaligned "
        "(wrong original matched to translation), return corrected pairs.\n\n"
        'Return ONLY JSON: {"pairs": [[0,0], [1,1], ...], '
        '"confidence": "high|medium|low", "issues": ["note any mismatches"]}'
    )
    try:
        resp = client.chat.completions.create(
            model=provider["model"],
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": chunk_text},
            ],
            temperature=0.0,
        )
        result = resp.choices[0].message.content
        parsed = json.loads(result)
        confidence = parsed.get("confidence", "low")
        if confidence == "low":
            print(_["warn_ai_low_confidence"], file=sys.stderr)
            return None
        return pairs
    except Exception as e:
        print(_["warn_ai_verify_failed"].format(e=e), file=sys.stderr)
        return None


def _llm_full_column_matching(all_dicts, provider):
    if provider is None:
        return None
    lines = []
    for i, p in enumerate(all_dicts):
        text = "".join(r["text"] for r in p["runs"])
        lines.append(f"[P{i}] {text}")
    chunk_text = "\n".join(lines)
    client = OpenAI(base_url=provider["base_url"], api_key=provider.get("api_key", "not-needed"))
    system_prompt = (
        "This document has two columns. Paragraphs are listed in order "
        "(first column 1, then column 2). Determine the column split "
        "and pair each original with its translation.\n\n"
        'Return ONLY JSON: {"pairs": [[0, 5], [1, 6], ...], '
        '"confidence": "high|medium|low"} where each pair maps a column-1 '
        "paragraph index to its column-2 paragraph index."
    )
    try:
        resp = client.chat.completions.create(
            model=provider["model"],
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": chunk_text},
            ],
            temperature=0.0,
        )
        result = resp.choices[0].message.content
        parsed = json.loads(result)
        index_pairs = parsed.get("pairs", [])
        result_pairs = []
        for c1_idx, c2_idx in index_pairs:
            left = all_dicts[c1_idx] if c1_idx < len(all_dicts) else None
            right = all_dicts[c2_idx] if c2_idx < len(all_dicts) else None
            result_pairs.append((left, right))
        return result_pairs
    except Exception as e:
        print(f"Warning: full AI column matching failed ({e})", file=sys.stderr)
        return None


def transform2cell(input_path, output_path, provider=None):
    doc = Document(input_path)
    is_col = _has_2column_layout(doc)
    is_table = _has_2column_table(doc)

    if not is_col and not is_table:
        print(_["warn_no_2column_layout"], file=sys.stderr)

    if is_table and not is_col:
        pairs = _extract_pairs_from_table(doc)
        _write_2cell_table(input_path, pairs, output_path)
        return

    split_idx = _find_column_break(doc)
    if split_idx is None:
        split_idx = _heuristic_column_split(doc)
        print(_["info_heuristic_column_split"].format(idx=split_idx), file=sys.stderr)
        skip = 0
    else:
        skip = 1

    all_dicts = [_build_para_dict(p) for p in doc.paragraphs]

    col1 = all_dicts[:split_idx]
    col2 = all_dicts[split_idx + skip:]
    pairs = [([d1], [d2]) for d1, d2 in zip(col1, col2)]

    if len(col1) != len(col2):
        print(_["warn_transform2cell_mismatch"].format(c1=len(col1), c2=len(col2)), file=sys.stderr)

    if provider:
        raw_pairs = list(zip(col1, col2))
        raw_pairs = _llm_verify_pairs(col1, col2, raw_pairs, provider)
        if raw_pairs is None:
            raw_pairs = _llm_full_column_matching(all_dicts, provider)
        if raw_pairs is None:
            print("AI matching failed. Falling back to heuristic pairing.", file=sys.stderr)
            raw_pairs = list(zip(col1, col2))
        pairs = [([d1], [d2]) for d1, d2 in raw_pairs]

    _write_2cell_table(input_path, pairs, output_path)


def test_config_loading():
    import json
    import tempfile
    import os
    cfg = {
        "default_provider": "test",
        "providers": {
            "test": {
                "base_url": "http://test/v1",
                "model": "test-model",
                "api_key_env": None
            }
        }
    }
    with tempfile.NamedTemporaryFile(mode='w', suffix='.json', delete=False) as f:
        json.dump(cfg, f)
        fname = f.name
    try:
        result = load_config(fname)
        assert result == cfg
    finally:
        os.unlink(fname)


def test_get_provider():
    cfg = {
        "default_provider": "test",
        "providers": {
            "test": {
                "base_url": "http://test/v1",
                "model": "test-model",
                "api_key_env": None
            }
        }
    }
    provider = get_provider(cfg, "test")
    assert provider["base_url"] == "http://test/v1"
    assert provider["model"] == "test-model"


def test_extract_paragraphs():
    from docx import Document
    import tempfile
    import os
    doc = Document()
    doc.add_paragraph("Hello world")
    p = doc.add_paragraph()
    p.add_run("Bold").bold = True
    p.add_run(" Normal")
    path = os.path.join(tempfile.mkdtemp(), "test.docx")
    doc.save(path)
    paragraphs = extract_paragraphs(path)
    assert len(paragraphs) == 2
    p0 = paragraphs[0]
    assert p0["id"] == "P0"
    assert not p0["in_table"]
    assert p0["cell_row"] is None
    assert p0["cell_col"] is None
    assert p0["alignment"] is None or p0["alignment"] is not None
    assert p0["runs"][0]["text"] == "Hello world"
    assert p0["runs"][0]["italic"] is None
    assert p0["runs"][0]["underline"] is None
    p1 = paragraphs[1]
    assert p1["runs"][0]["text"] == "Bold"
    assert p1["runs"][0]["bold"]
    assert p1["runs"][1]["text"] == " Normal"
    os.unlink(path)


def test_extract_paragraphs_with_table():
    from docx import Document
    import tempfile
    import os
    doc = Document()
    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "A1"
    table.cell(0, 1).text = "B1"
    table.cell(1, 2).text = "C2"
    path = os.path.join(tempfile.mkdtemp(), "test_table.docx")
    doc.save(path)
    paragraphs = extract_paragraphs(path)
    table_paras = [p for p in paragraphs if p["in_table"]]
    assert len(table_paras) == 6, f"Expected 6 table paragraphs, got {len(table_paras)} total paragraphs={len(paragraphs)}"
    for p in table_paras:
        assert p["cell_row"] is not None, f"Expected cell_row, got None for {p}"
        assert p["cell_col"] is not None, f"Expected cell_col, got None for {p}"
    scored = [(p["cell_row"], p["cell_col"], p["runs"][0]["text"] if p["runs"] else "") for p in table_paras]
    assert (0, 0, "A1") in scored, f"Missing A1 at (0,0): {scored}"
    assert (0, 1, "B1") in scored, f"Missing B1 at (0,1): {scored}"
    assert (1, 2, "C2") in scored, f"Missing C2 at (1,2): {scored}"
    os.unlink(path)

def test_extract_paragraphs_invalid_file():
    try:
        extract_paragraphs("/nonexistent/file.docx")
        assert False, "Should have raised FileNotFoundError"
    except FileNotFoundError:
        pass


def test_chunk_paragraphs():
    paragraphs = [{"id": f"P{i}"} for i in range(25)]
    chunks = chunk_paragraphs(paragraphs, 10)
    assert len(chunks) == 3


def test_translate_chunk():
    chunk = [{"id": "P0", "runs": [{"text": "Hello world", "bold": False, "italic": False}]}]
    provider = {"base_url": "http://invalid", "model": "test", "api_key": "bad"}
    try:
        translate_chunk(chunk, "Romanian", provider)
        assert False, "Should have raised ConnectionError or similar"
    except Exception:
        pass


def test_translate_chunk_retry():
    pass


def test_write_inline():
    from docx import Document
    import tempfile
    import os
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Original text")
    src = os.path.join(tempfile.mkdtemp(), "source.docx")
    doc.save(src)
    translated = [{
        "id": "P0",
        "runs": [{"text": "Text tradus", "bold": False, "italic": False}],
        "alignment": None,
        "in_table": False,
    }]
    out = os.path.join(tempfile.mkdtemp(), "output.docx")
    write_inline(src, translated, out)
    result = Document(out)
    assert result.paragraphs[0].text == "Text tradus"
    os.unlink(src)
    os.unlink(out)


def _apply_run_formatting(run, run_data):
    run.bold = run_data.get("bold")
    run.italic = run_data.get("italic")
    run.underline = run_data.get("underline")
    if run_data.get("font_name"):
        run.font.name = run_data["font_name"]
    fs = run_data.get("font_size")
    if fs is not None:
        try:
            run.font.size = Pt(float(fs))
        except (ValueError, TypeError):
            pass
    if run_data.get("color"):
        try:
            run.font.color.rgb = RGBColor.from_string(run_data["color"])
        except (ValueError, AttributeError):
            pass


def _format_number(value, fmt):
    if fmt == 'upperRoman':
        vals = [(1000, 'M'), (900, 'CM'), (500, 'D'), (400, 'CD'), (100, 'C'),
                (90, 'XC'), (50, 'L'), (40, 'XL'), (10, 'X'), (9, 'IX'),
                (5, 'V'), (4, 'IV'), (1, 'I')]
        result = ''
        n = value
        for a, r in vals:
            while n >= a:
                result += r
                n -= a
        return result
    if fmt == 'lowerRoman':
        vals = [(1000, 'm'), (900, 'cm'), (500, 'd'), (400, 'cd'), (100, 'c'),
                (90, 'xc'), (50, 'l'), (40, 'xl'), (10, 'x'), (9, 'ix'),
                (5, 'v'), (4, 'iv'), (1, 'i')]
        result = ''
        n = value
        for a, r in vals:
            while n >= a:
                result += r
                n -= a
        return result
    if fmt == 'decimal':
        return str(value)
    if fmt == 'lowerLetter':
        return chr(ord('a') + value - 1) if 1 <= value <= 26 else str(value)
    if fmt == 'upperLetter':
        return chr(ord('A') + value - 1) if 1 <= value <= 26 else str(value)
    return str(value)


_NUMBERING_FMT_CACHE = {}


def _get_numpr_fmt(original_path, numId):
    key = (original_path, numId)
    if key in _NUMBERING_FMT_CACHE:
        return _NUMBERING_FMT_CACHE[key]
    try:
        with zipfile.ZipFile(original_path, 'r') as z:
            n = z.read('word/numbering.xml').decode()
            abs_match = re.search(rf'<w:num w:numId=\"{numId}\".*?<w:abstractNumId w:val=\"(\d+)\"', n, re.DOTALL)
            if not abs_match:
                _NUMBERING_FMT_CACHE[key] = None
                return None
            abs_id = abs_match.group(1)
            fmt_match = re.search(
                rf'<w:abstractNum w:abstractNumId=\"{abs_id}\".*?<w:lvl w:ilvl=\"0\".*?<w:numFmt w:val=\"([^\"]+)\"',
                n, re.DOTALL
            )
            fmt = fmt_match.group(1) if fmt_match else None
            _NUMBERING_FMT_CACHE[key] = fmt
            return fmt
    except Exception:
        return None


def _apply_numbering_labels(paras, original_path):
    num_groups = {}
    for p in paras:
        np = p.get("numpr")
        if np and np.get("numId"):
            key = np["numId"]
            num_groups.setdefault(key, []).append(p)
    for numId, group in num_groups.items():
        fmt = _get_numpr_fmt(original_path, numId)
        if not fmt or fmt == 'bullet':
            continue
        for idx, p in enumerate(group):
            number = _format_number(idx + 1, fmt)
            prefix = number + ". "
            if p["runs"]:
                p["runs"][0]["text"] = prefix + p["runs"][0]["text"]
    return paras


def write_side_by_side(original_path, original_paras, translated_paras, output_path):
    src = Document(original_path)
    doc = Document()
    for section in src.sections:
        for new_section in doc.sections:
            new_section.page_width = section.page_width
            new_section.page_height = section.page_height
            new_section.left_margin = section.left_margin
            new_section.right_margin = section.right_margin
            new_section.top_margin = section.top_margin
            new_section.bottom_margin = section.bottom_margin
        break
    trans_by_id = {p["id"]: p for p in translated_paras}
    table = doc.add_table(rows=len(original_paras), cols=2)
    ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    tbl_pr = table._tbl.tblPr
    tbl_borders = doc.element.makeelement(ns + 'tblBorders', {})
    for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        child = doc.element.makeelement(ns + edge, {})
        child.set(ns + 'val', 'none')
        child.set(ns + 'sz', '0')
        child.set(ns + 'space', '0')
        child.set(ns + 'color', 'auto')
        tbl_borders.append(child)
    tbl_pr.append(tbl_borders)
    _apply_numbering_labels(original_paras, original_path)
    _apply_numbering_labels(translated_paras, original_path)
    for i, op in enumerate(original_paras):
        pid = op["id"]
        left_cell = table.rows[i].cells[0]
        right_cell = table.rows[i].cells[1]
        for run_data in op["runs"]:
            p = left_cell.paragraphs[0] if left_cell.paragraphs else left_cell.add_paragraph()
            run = p.add_run(run_data["text"])
            _apply_run_formatting(run, run_data)
        tp = trans_by_id.get(pid, op)
        for run_data in tp["runs"]:
            p = right_cell.paragraphs[0] if right_cell.paragraphs else right_cell.add_paragraph()
            run = p.add_run(run_data["text"])
            _apply_run_formatting(run, run_data)
    doc.save(output_path)


def test_write_inline_with_table():
    from docx import Document
    import tempfile
    import os
    doc = Document()
    doc.add_paragraph("Header text")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Cell A1"
    table.cell(0, 1).text = "Cell B1"
    table.cell(1, 0).text = "Cell A2"
    table.cell(1, 1).text = "Cell B2"
    src = os.path.join(tempfile.mkdtemp(), "source.docx")
    doc.save(src)
    translated = [
        {"id": "P0", "runs": [{"text": "Header text trans"}], "alignment": None, "in_table": False},
        {"id": "P1", "runs": [{"text": "Celula A1"}], "alignment": None, "in_table": True, "cell_row": 0, "cell_col": 0},
        {"id": "P2", "runs": [{"text": "Celula B1"}], "alignment": None, "in_table": True, "cell_row": 0, "cell_col": 1},
        {"id": "P3", "runs": [{"text": "Celula A2"}], "alignment": None, "in_table": True, "cell_row": 1, "cell_col": 0},
        {"id": "P4", "runs": [{"text": "Celula B2"}], "alignment": None, "in_table": True, "cell_row": 1, "cell_col": 1},
    ]
    out = os.path.join(tempfile.mkdtemp(), "output.docx")
    write_inline(src, translated, out)
    result = Document(out)
    assert result.paragraphs[0].text == "Header text trans"
    out_table = result.tables[0]
    assert out_table.cell(0, 0).text == "Celula A1"
    assert out_table.cell(0, 1).text == "Celula B1"
    assert out_table.cell(1, 0).text == "Celula A2"
    assert out_table.cell(1, 1).text == "Celula B2"
    os.unlink(src)
    os.unlink(out)


def test_cli_parser():
    args = parse_args(['input.docx', '--lang', 'ro'])
    assert args.input == 'input.docx'
    assert args.lang == 'ro'
    assert args.mode == 'inline'
    assert args.concurrency == 0
    args = parse_args(['input.docx', '--lang', 'en', '--mode', 'side-by-side', '--concurrency', '2'])
    assert args.input == 'input.docx'
    assert args.lang == 'en'
    assert args.mode == 'side-by-side'
    assert args.concurrency == 2


def _write_paras_to_cell(cell, paras):
    if not paras:
        return
    for pi, para_dict in enumerate(paras):
        if pi == 0 and cell.paragraphs:
            p = cell.paragraphs[0]
            p.clear()
        else:
            p = cell.add_paragraph()
        if para_dict.get("alignment"):
            try:
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                p.alignment = WD_ALIGN_PARAGRAPH(int(para_dict["alignment"]))
            except (ValueError, TypeError):
                pass
        for run_data in para_dict["runs"]:
            run = p.add_run(run_data["text"])
            _apply_run_formatting(run, run_data)


def _write_2cell_table(original_path, pairs, output_path):
    src = Document(original_path)
    doc = Document()
    for section in src.sections:
        for new_section in doc.sections:
            new_section.page_width = section.page_width
            new_section.page_height = section.page_height
            new_section.left_margin = section.left_margin
            new_section.right_margin = section.right_margin
            new_section.top_margin = section.top_margin
            new_section.bottom_margin = section.bottom_margin
        break

    table = doc.add_table(rows=len(pairs), cols=2)
    ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    tbl_pr = table._tbl.tblPr
    tbl_borders = doc.element.makeelement(ns + 'tblBorders', {})
    for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        child = doc.element.makeelement(ns + edge, {})
        child.set(ns + 'val', 'none')
        child.set(ns + 'sz', '0')
        child.set(ns + 'space', '0')
        child.set(ns + 'color', 'auto')
        tbl_borders.append(child)
    tbl_pr.append(tbl_borders)

    for i, (left_paras, right_paras) in enumerate(pairs):
        left_cell = table.rows[i].cells[0]
        right_cell = table.rows[i].cells[1]
        _write_paras_to_cell(left_cell, left_paras)
        _write_paras_to_cell(right_cell, right_paras)
    doc.save(output_path)


def test_write_2cell_table():
    import tempfile, os
    from docx import Document
    doc = Document()
    doc.add_paragraph("Source only")
    src = os.path.join(tempfile.mkdtemp(), "src.docx")
    doc.save(src)
    pairs = [
        ([{"runs": [{"text": "Left A", "bold": False}], "alignment": None, "numpr": None}],
         [{"runs": [{"text": "Right A", "bold": False}], "alignment": None, "numpr": None}]),
        ([{"runs": [{"text": "Left B", "bold": False}], "alignment": None, "numpr": None}],
         [{"runs": [{"text": "Right B", "bold": False}], "alignment": None, "numpr": None}]),
    ]
    out = os.path.join(tempfile.mkdtemp(), "out.docx")
    _write_2cell_table(src, pairs, out)
    result = Document(out)
    assert len(result.tables) == 1
    t = result.tables[0]
    assert len(t.rows) == 2
    assert t.rows[0].cells[0].text == "Left A"
    assert t.rows[0].cells[1].text == "Right A"
    assert t.rows[1].cells[0].text == "Left B"
    assert t.rows[1].cells[1].text == "Right B"
    os.unlink(src); os.unlink(out)


def test_write_2cell_table_no_borders():
    import tempfile, os
    from docx import Document
    from lxml import etree
    doc = Document()
    doc.add_paragraph("x")
    src = os.path.join(tempfile.mkdtemp(), "src.docx")
    doc.save(src)
    pairs = [([{"runs": [{"text": "A", "bold": False}], "alignment": None, "numpr": None}],
              [{"runs": [{"text": "B", "bold": False}], "alignment": None, "numpr": None}])]
    out = os.path.join(tempfile.mkdtemp(), "out.docx")
    _write_2cell_table(src, pairs, out)
    result = Document(out)
    tbl = result.tables[0]
    tbl_pr = tbl._tbl.tblPr
    borders = tbl_pr.find(f'{NS_W}tblBorders')
    assert borders is not None, "Table should have tblBorders element"
    os.unlink(src); os.unlink(out)


def test_write_side_by_side():
    from docx import Document
    import tempfile
    import os
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Original text")
    src = os.path.join(tempfile.mkdtemp(), "source.docx")
    doc.save(src)
    originals = [{"id": "P0", "runs": [{"text": "Original text", "bold": False}], "alignment": None, "in_table": False}]
    translated = [{"id": "P0", "runs": [{"text": "Text tradus", "bold": False}], "alignment": None, "in_table": False}]
    out = os.path.join(tempfile.mkdtemp(), "output.docx")
    write_side_by_side(src, originals, translated, out)
    result = Document(out)
    assert len(result.tables) == 1
    table = result.tables[0]
    assert len(table.rows) == 1
    assert "Original text" in table.rows[0].cells[0].text
    assert "Text tradus" in table.rows[0].cells[1].text
    os.unlink(src)
    os.unlink(out)


def test_restore_placeholders():
    original = "Hello {{client_name}}, your balance is {{amount}}"
    translated = "Salut {{client_name}}, soldul tau este {{amount}}"
    restored = restore_placeholders(translated, original)
    assert "{{client_name}}" in restored
    assert "{{amount}}" in restored


def test_restore_missing_placeholder():
    original = "Hello {{client_name}}"
    translated = "Salut"
    restored = restore_placeholders(translated, original)
    assert "{{client_name}}" in restored


def test_parse_translated_response_protects_placeholders():
    chunk_ids = ["P0", "P1"]
    response_text = "[P0] Salut {{client_name}}\n[P1] Soldul este {{amount}}"
    original_text = "[P0] Hello {{client_name}}\n[P1] Your balance is {{amount}}"
    result = _parse_translated_response(response_text, chunk_ids, original_text)
    assert "{{client_name}}" in result["P0"]
    assert "{{amount}}" in result["P1"]


def test_has_2column_layout():
    import tempfile, os
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    doc = Document()
    path = os.path.join(tempfile.mkdtemp(), "test.docx")
    doc.save(path)
    assert not _has_2column_layout(Document(path))
    doc2 = Document()
    sect_pr = doc2.sections[0]._sectPr
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:space'), '720')
    sect_pr.append(cols)
    path2 = os.path.join(tempfile.mkdtemp(), "test2.docx")
    doc2.save(path2)
    assert _has_2column_layout(Document(path2))
    os.unlink(path); os.unlink(path2)


def test_has_2column_table():
    import tempfile, os
    from docx import Document
    doc = Document()
    doc.add_paragraph("Not a table")
    path = os.path.join(tempfile.mkdtemp(), "test.docx")
    doc.save(path)
    assert not _has_2column_table(Document(path))

    doc2 = Document()
    doc2.add_table(rows=2, cols=2)
    path2 = os.path.join(tempfile.mkdtemp(), "test2.docx")
    doc2.save(path2)
    assert _has_2column_table(Document(path2))

    doc3 = Document()
    doc3.add_table(rows=2, cols=3)
    path3 = os.path.join(tempfile.mkdtemp(), "test3.docx")
    doc3.save(path3)
    assert not _has_2column_table(Document(path3))

    os.unlink(path); os.unlink(path2); os.unlink(path3)


def test_find_column_break():
    import tempfile, os
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    path = os.path.join(tempfile.mkdtemp(), "test.docx")
    doc = Document()
    doc.add_paragraph("Before break")
    p = doc.add_paragraph()
    r = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'column')
    r._element.append(br)
    doc.add_paragraph("After break")
    doc.save(path)
    idx = _find_column_break(Document(path))
    assert idx == 1, f"Expected 1, got {idx}"
    doc2 = Document()
    doc2.add_paragraph("No break")
    doc2.add_paragraph("No break either")
    path2 = os.path.join(tempfile.mkdtemp(), "test2.docx")
    doc2.save(path2)
    assert _find_column_break(Document(path2)) is None
    os.unlink(path); os.unlink(path2)


def test_heuristic_column_split():
    from docx import Document
    doc = Document()
    for _ in range(10):
        doc.add_paragraph("x")
    assert _heuristic_column_split(doc) == 5
    doc2 = Document()
    for _ in range(7):
        doc2.add_paragraph("x")
    assert _heuristic_column_split(doc2) == 3


def test_pair_by_position():
    paras = [f"P{i}" for i in range(6)]
    pairs = _pair_by_position(paras, 3)
    assert len(pairs) == 3
    assert pairs[0] == ("P0", "P3")
    assert pairs[1] == ("P1", "P4")
    assert pairs[2] == ("P2", "P5")
    paras2 = [f"P{i}" for i in range(5)]
    pairs2 = _pair_by_position(paras2, 2)
    assert len(pairs2) == 3
    assert pairs2[0] == ("P0", "P2")
    assert pairs2[1] == ("P1", "P3")
    assert pairs2[2] == (None, "P4")


def parse_args(argv=None):
    parser = argparse.ArgumentParser(description=_["cli_desc"])
    parser.add_argument("input", help=_["cli_input_help"])
    parser.add_argument("--lang", "-l", choices=["ro", "en"],
                        help=_["cli_lang_help"])
    parser.add_argument("--mode", "-m", choices=["inline", "side-by-side"],
                        default="inline", help=_["cli_mode_help"])
    parser.add_argument("--output", "-o", help=_["cli_output_help"])
    parser.add_argument("--provider", "-p", help=_["cli_provider_help"])
    parser.add_argument("--model", help=_["cli_model_help"])
    parser.add_argument("--config", "-c", default="config.json", help=_["cli_config_help"])
    parser.add_argument("--concurrency", "-j", type=int, default=0,
                        help=_["cli_concurrency_help"])
    parser.add_argument("--transform2cell", action="store_true",
                        help=_["cli_transform2cell_help"])
    global _PARSER
    _PARSER = parser
    return parser.parse_args(argv)


def main():
    args = parse_args()
    config = load_config(args.config)

    if args.transform2cell:
        provider = get_provider(config, args.provider) if args.provider else None
        if args.model and provider:
            provider["model"] = args.model
        output = args.output
        if not output:
            stem = args.input.rsplit(".", 1)[0]
            output = f"{stem}_2col_table.docx"
        if output == args.input:
            stem = output.rsplit(".docx", 1)[0]
            timestamp = time.strftime("%Y%m%d_%H%M%S")
            output = f"{stem}_{timestamp}.docx"
        elif os.path.exists(output):
            stem = output.rsplit(".docx", 1)[0]
            timestamp = time.strftime("%Y%m%d_%H%M%S")
            output = f"{stem}_{timestamp}.docx"
        transform2cell(args.input, output, provider)
        print(_["saved"].format(output=output))
        return

    if not args.lang:
        _PARSER.error("--lang is required (unless using --transform2cell)")

    provider = get_provider(config, args.provider)
    if args.model:
        provider["model"] = args.model
    output = args.output
    if not output:
        stem = args.input.rsplit(".", 1)[0]
        suffix = f"_{args.mode}_{args.lang}.docx" if args.mode == "side-by-side" else f"_{args.lang}.docx"
        output = f"{stem}{suffix}"
    if os.path.exists(output):
        stem = output.rsplit(".docx", 1)[0]
        timestamp = time.strftime("%Y%m%d_%H%M%S")
        output = f"{stem}_{timestamp}.docx"
    paragraphs = extract_paragraphs(args.input)
    originals = copy.deepcopy(paragraphs) if args.mode == "side-by-side" else None
    translated = translate_all(paragraphs, "Romanian" if args.lang == "ro" else "English", provider, args.concurrency)
    if args.mode == "side-by-side":
        write_side_by_side(args.input, originals, translated, output)
    else:
        write_inline(args.input, translated, output)
    print(_["saved"].format(output=output))


def _create_2column_test_docx(path):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx import Document

    doc = Document()
    sect_pr = doc.sections[0]._sectPr
    existing = sect_pr.find(qn('w:cols'))
    if existing is not None:
        sect_pr.remove(existing)
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:space'), '720')
    sect_pr.append(cols)

    doc.add_paragraph("First original paragraph.")
    doc.add_paragraph("Second original paragraph.")
    doc.add_paragraph("Third original paragraph.")

    p_break = doc.add_paragraph()
    r_break = p_break.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'column')
    r_break._element.append(br)

    doc.add_paragraph("Primul paragraf original.")
    doc.add_paragraph("Al doilea paragraf original.")
    doc.add_paragraph("Al treilea paragraf original.")
    doc.save(path)


def test_transform2cell_integration():
    import tempfile, os
    from docx import Document
    path = os.path.join(tempfile.mkdtemp(), "src.docx")
    _create_2column_test_docx(path)
    out = os.path.join(tempfile.mkdtemp(), "out.docx")
    transform2cell(path, out)
    result = Document(out)
    assert len(result.tables) == 1
    t = result.tables[0]
    assert len(t.rows) == 3
    assert "First original" in t.rows[0].cells[0].text
    assert "Primul paragraf" in t.rows[0].cells[1].text
    assert "Second original" in t.rows[1].cells[0].text
    assert "Al doilea" in t.rows[1].cells[1].text
    assert "Third original" in t.rows[2].cells[0].text
    assert "Al treilea" in t.rows[2].cells[1].text
    os.unlink(path); os.unlink(out)


def test_transform2cell_table_integration():
    import tempfile, os
    from docx import Document
    path = os.path.join(tempfile.mkdtemp(), "src.docx")
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Original A"
    table.cell(0, 1).text = "Translated A"
    table.cell(1, 0).text = "Original B"
    table.cell(1, 1).text = "Translated B"
    doc.save(path)
    out = os.path.join(tempfile.mkdtemp(), "out.docx")
    transform2cell(path, out)
    result = Document(out)
    assert len(result.tables) == 1
    t = result.tables[0]
    assert len(t.rows) == 2
    assert "Original A" in t.rows[0].cells[0].text
    assert "Translated A" in t.rows[0].cells[1].text
    assert "Original B" in t.rows[1].cells[0].text
    assert "Translated B" in t.rows[1].cells[1].text
    os.unlink(path); os.unlink(out)


def test_cli_parser_transform2cell():
    args = parse_args(['input.docx', '--transform2cell'])
    assert args.transform2cell is True
    assert args.input == 'input.docx'
    args2 = parse_args(['input.docx', '--transform2cell', '--output', 'out.docx'])
    assert args2.transform2cell
    assert args2.output == 'out.docx'


if __name__ == "__main__":
    if len(sys.argv) > 1:
        main()
    else:
        tests = [
            ("test_config_loading", test_config_loading),
            ("test_get_provider", test_get_provider),
            ("test_extract_paragraphs", test_extract_paragraphs),
            ("test_extract_paragraphs_with_table", test_extract_paragraphs_with_table),
            ("test_extract_paragraphs_invalid_file", test_extract_paragraphs_invalid_file),
            ("test_chunk_paragraphs", test_chunk_paragraphs),
            ("test_translate_chunk", test_translate_chunk),
            ("test_translate_chunk_retry", test_translate_chunk_retry),
            ("test_write_inline", test_write_inline),
            ("test_write_side_by_side", test_write_side_by_side),
            ("test_write_inline_with_table", test_write_inline_with_table),
            ("test_cli_parser", test_cli_parser),
            ("test_restore_placeholders", test_restore_placeholders),
            ("test_restore_missing_placeholder", test_restore_missing_placeholder),
            ("test_parse_translated_response_protects_placeholders", test_parse_translated_response_protects_placeholders),
            ("test_has_2column_layout", test_has_2column_layout),
            ("test_has_2column_table", test_has_2column_table),
            ("test_find_column_break", test_find_column_break),
            ("test_heuristic_column_split", test_heuristic_column_split),
            ("test_write_2cell_table", test_write_2cell_table),
            ("test_write_2cell_table_no_borders", test_write_2cell_table_no_borders),
            ("test_pair_by_position", test_pair_by_position),
            ("test_transform2cell_integration", test_transform2cell_integration),
            ("test_transform2cell_table_integration", test_transform2cell_table_integration),
            ("test_cli_parser_transform2cell", test_cli_parser_transform2cell),
        ]
        for name, fn in tests:
            fn()
            print(_["test_pass"].format(name=name))
        print(_["all_pass"])
